from io import BytesIO

from docx import Document
from fastapi import HTTPException

# ==========================================================
# Limits
# ==========================================================

MAX_PARAGRAPHS = 500
MAX_TABLES = 100


# ==========================================================
# Open DOCX
# ==========================================================

def open_docx(data: bytes):
    """
    Open DOCX document from bytes.
    """

    try:
        return Document(BytesIO(data))

    except Exception:
        raise HTTPException(
            status_code=400,
            detail="Invalid or corrupted DOCX file.",
        )


# ==========================================================
# Paragraph Extraction
# ==========================================================

def extract_paragraphs(document):
    """
    Extract all paragraphs.
    """

    paragraphs = []

    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:
            paragraphs.append(text)

    return paragraphs


# ==========================================================
# Table Extraction
# ==========================================================

def extract_tables(document):
    """
    Extract text from DOCX tables.
    """

    rows = []

    for table in document.tables:

        for row in table.rows:

            values = []

            for cell in row.cells:

                value = cell.text.strip()

                if value:
                    values.append(value)

            if values:
                rows.append(" | ".join(values))

    return rows


# ==========================================================
# Header Extraction
# ==========================================================

def extract_headers(document):
    """
    Extract text from document headers.
    """

    headers = []

    for section in document.sections:

        for paragraph in section.header.paragraphs:

            text = paragraph.text.strip()

            if text:
                headers.append(text)

    return headers


# ==========================================================
# Footer Extraction
# ==========================================================

def extract_footers(document):
    """
    Extract text from document footers.
    """

    footers = []

    for section in document.sections:

        for paragraph in section.footer.paragraphs:

            text = paragraph.text.strip()

            if text:
                footers.append(text)

    return footers


# ==========================================================
# Hyperlink Extraction
# ==========================================================

def extract_hyperlinks(document):
    """
    Extract hyperlinks from relationships.
    """

    links = []
    seen = set()

    try:

        for relation in document.part.rels.values():

            if (
                relation.reltype.endswith("/hyperlink")
                and relation.target_ref
            ):

                url = relation.target_ref.strip()

                if url and url not in seen:
                    seen.add(url)
                    links.append(url)

    except Exception:
        pass

    return links


# ==========================================================
# Merge Everything
# ==========================================================

def merge_document_content(
    headers,
    paragraphs,
    tables,
    footers,
    hyperlinks,
):
    """
    Merge extracted content.
    """

    content = []

    if headers:
        content.extend(headers)
        content.append("")

    if paragraphs:
        content.extend(paragraphs)
        content.append("")

    if tables:
        content.extend(tables)
        content.append("")

    if footers:
        content.extend(footers)
        content.append("")

    if hyperlinks:
        content.append("")
        content.extend(hyperlinks)

    return "\n".join(content).strip()


# ==========================================================
# Parse DOCX
# ==========================================================

def parse_docx(data: bytes) -> str:
    """
    Parse DOCX into plain text.
    """

    document = open_docx(data)

    headers = extract_headers(document)

    paragraphs = extract_paragraphs(document)

    tables = extract_tables(document)

    footers = extract_footers(document)

    hyperlinks = extract_hyperlinks(document)

    return merge_document_content(
        headers,
        paragraphs,
        tables,
        footers,
        hyperlinks,
    )


# ==========================================================
# Statistics
# ==========================================================

def extract_docx_statistics(document):
    """
    Extract document statistics.
    """

    paragraph_count = len(document.paragraphs)
    table_count = len(document.tables)
    section_count = len(document.sections)

    word_count = sum(
        len(paragraph.text.split())
        for paragraph in document.paragraphs
    )

    return {
        "paragraphs": paragraph_count,
        "tables": table_count,
        "sections": section_count,
        "words": word_count,
    }


# ==========================================================
# Metadata
# ==========================================================

def extract_docx_metadata(document):
    """
    Extract DOCX metadata.
    """

    props = document.core_properties

    return {
        "title": props.title or None,
        "author": props.author or None,
        "subject": props.subject or None,
        "keywords": props.keywords or None,
        "comments": props.comments or None,
        "category": props.category or None,
        "created": (
            props.created.isoformat()
            if props.created
            else None
        ),
        "modified": (
            props.modified.isoformat()
            if props.modified
            else None
        ),
    }


# ==========================================================
# Validation
# ==========================================================

def validate_docx_text(text: str):
    """
    Validate parsed text.
    """

    if not text:
        raise HTTPException(
            status_code=400,
            detail="No readable text found in DOCX.",
        )

    if len(text.split()) < 10:
        raise HTTPException(
            status_code=400,
            detail="DOCX contains too little readable text.",
        )


# ==========================================================
# Full Analysis
# ==========================================================

def analyze_docx(data: bytes):
    """
    Parse DOCX and return analysis.
    """

    document = open_docx(data)

    headers = extract_headers(document)
    paragraphs = extract_paragraphs(document)
    tables = extract_tables(document)
    footers = extract_footers(document)
    hyperlinks = extract_hyperlinks(document)

    text = merge_document_content(
        headers,
        paragraphs,
        tables,
        footers,
        hyperlinks,
    )

    validate_docx_text(text)

    return {
        "text": text,
        "metadata": extract_docx_metadata(document),
        "statistics": extract_docx_statistics(document),
        "links": hyperlinks,
    }


# ==========================================================
# Public API
# ==========================================================

def extract_docx_text(data: bytes) -> str:
    """
    Public helper.
    """

    return analyze_docx(data)["text"]


# ==========================================================
# Debug
# ==========================================================

if __name__ == "__main__":
    print("docx_parser.py loaded successfully.")